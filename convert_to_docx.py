#!/usr/bin/env python3
"""
ONYX Security Architecture Document Converter
Converts the Markdown security document to a professional DOCX format
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add a page break to the document"""
    paragraph = doc.add_paragraph()
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

def create_custom_styles(doc):
    """Create custom styles for the document"""
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 32, 96)  # Dark blue
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Header 1 style
    h1_style = styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 32, 96)
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    h1_style.paragraph_format.keep_with_next = True
    
    # Header 2 style
    h2_style = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(16)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 70, 140)
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(4)
    
    # Header 3 style
    h3_style = styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
    h3_font = h3_style.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(14)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(0, 100, 180)
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(3)
    
    # Code style
    code_style = styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)
    
    # Table header style
    table_header_style = styles.add_style('CustomTableHeader', WD_STYLE_TYPE.PARAGRAPH)
    table_header_font = table_header_style.font
    table_header_font.name = 'Calibri'
    table_header_font.size = Pt(11)
    table_header_font.bold = True
    table_header_font.color.rgb = RGBColor(255, 255, 255)
    table_header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_cover_page(doc):
    """Add a professional cover page"""
    # Title
    title = doc.add_paragraph()
    title.style = 'CustomTitle'
    title_run = title.add_run('ONYX Platform\nSecurity Architecture Document')
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Comprehensive Security Framework\nfor Building Assessment Platform')
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Add some space
    for _ in range(3):
        doc.add_paragraph()
    
    # Platform information
    platform_info = doc.add_paragraph()
    platform_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    platform_run = platform_info.add_run('Platform: onyxreport.com\nSecurity Level: Enterprise Grade')
    platform_run.font.size = Pt(14)
    platform_run.font.bold = True
    
    # Add more space
    for _ in range(4):
        doc.add_paragraph()
    
    # Document metadata
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata_text = """Document Version: 1.0
Last Updated: December 2024
Classification: Confidential
Prepared by: ONYX Security Team"""
    
    metadata_run = metadata.add_run(metadata_text)
    metadata_run.font.size = Pt(12)
    metadata_run.font.color.rgb = RGBColor(80, 80, 80)
    
    add_page_break(doc)

def process_table(doc, table_lines):
    """Process markdown table and convert to Word table"""
    if len(table_lines) < 2:
        return
    
    # Parse header and data rows
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    data_rows = []
    
    for line in table_lines[2:]:  # Skip separator line
        if '|' in line:
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(row) == len(headers):
                data_rows.append(row)
    
    if not data_rows:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set cell background color
        cell_xml = cell._tc
        cell_props = cell_xml.get_or_add_tcPr()
        cell_shading = OxmlElement('w:shd')
        cell_shading.set(qn('w:fill'), '2F5597')  # Blue background
        cell_props.append(cell_shading)
    
    # Add data rows
    for row_data in data_rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            cell = row.cells[i]
            cell.text = cell_data

def convert_markdown_to_docx(md_file_path, output_file_path):
    """Convert markdown file to formatted DOCX"""
    
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Create custom styles
    create_custom_styles(doc)
    
    # Add cover page
    add_cover_page(doc)
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    current_table_lines = []
    in_code_block = False
    code_language = ''
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line[3:].strip()
                i += 1
                continue
            else:
                in_code_block = False
                i += 1
                continue
        
        if in_code_block:
            para = doc.add_paragraph()
            para.style = 'CustomCode'
            run = para.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and line.strip():
            current_table_lines.append(line)
            i += 1
            continue
        elif current_table_lines:
            process_table(doc, current_table_lines)
            current_table_lines = []
            doc.add_paragraph()  # Add space after table
        
        # Handle headers
        if line.startswith('# '):
            para = doc.add_paragraph()
            para.style = 'CustomTitle'
            para.add_run(line[2:])
        elif line.startswith('## '):
            para = doc.add_paragraph()
            para.style = 'CustomH1'
            para.add_run(line[3:])
        elif line.startswith('### '):
            para = doc.add_paragraph()
            para.style = 'CustomH2'
            para.add_run(line[4:])
        elif line.startswith('#### '):
            para = doc.add_paragraph()
            para.style = 'CustomH3'
            para.add_run(line[5:])
        
        # Handle horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        
        # Handle bold text
        elif '**' in line:
            para = doc.add_paragraph()
            parts = line.split('**')
            for j, part in enumerate(parts):
                run = para.add_run(part)
                if j % 2 == 1:  # Odd indices are bold
                    run.font.bold = True
        
        # Handle bullet points
        elif line.startswith('- '):
            para = doc.add_paragraph()
            para.style = 'List Bullet'
            para.add_run(line[2:])
        
        # Handle numbered lists
        elif re.match(r'^\d+\. ', line):
            para = doc.add_paragraph()
            para.style = 'List Number'
            para.add_run(re.sub(r'^\d+\. ', '', line))
        
        # Handle empty lines
        elif not line.strip():
            if i > 0 and lines[i-1].strip():  # Only add if previous line had content
                doc.add_paragraph()
        
        # Regular paragraphs
        else:
            if line.strip():
                para = doc.add_paragraph()
                
                # Check for inline code
                if '`' in line:
                    parts = line.split('`')
                    for j, part in enumerate(parts):
                        run = para.add_run(part)
                        if j % 2 == 1:  # Odd indices are code
                            run.font.name = 'Consolas'
                            run.font.size = Pt(10)
                else:
                    para.add_run(line)
        
        i += 1
    
    # Process any remaining table
    if current_table_lines:
        process_table(doc, current_table_lines)
    
    # Save document
    doc.save(output_file_path)
    print(f"Document saved as: {output_file_path}")

def main():
    """Main function"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(script_dir, 'ONYX_Security_Architecture_Document.md')
    docx_file = os.path.join(script_dir, 'ONYX_Security_Architecture_Document.docx')
    
    if not os.path.exists(md_file):
        print(f"Error: Markdown file not found: {md_file}")
        return
    
    try:
        convert_markdown_to_docx(md_file, docx_file)
        print("✅ Conversion completed successfully!")
        print(f"📄 Professional DOCX document created: {docx_file}")
    except Exception as e:
        print(f"❌ Error during conversion: {str(e)}")

if __name__ == "__main__":
    main()